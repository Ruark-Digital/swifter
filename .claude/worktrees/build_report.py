from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Page title
title = doc.add_paragraph()
run = title.add_run("Product Update — Contract Management Platform")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

sub = doc.add_paragraph()
r = sub.add_run("Reporting period: 27 April 2026 – 11 May 2026")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x52, 0x5B, 0x6B)

doc.add_paragraph(
    "The team shipped twenty-four improvements across the contract lifecycle this "
    "period. Highlights below, grouped by area of impact."
)

sections = [
    (
        "Multi-Currency Support",
        [
            "Contracts and Master Service Agreements (MSAs) can now be created and managed in "
            "currencies other than the company's default. When an MSA is opened in a different "
            "currency, the system automatically captures an exchange rate so reporting stays "
            "consistent across the portfolio. All money fields — milestones, holdbacks, "
            "securities, invoices — render with the proper currency symbol and locale formatting."
        ],
    ),
    (
        "Compliance & Security",
        [
            "The compliance area now supports dark mode end-to-end.",
            "The Contract Manager's Approve / Reject controls are now hidden until the vendor has "
            "actually submitted the policy or security document — preventing premature decisions "
            "on empty submissions.",
            "Once a compliance decision is in flight, all related action buttons across the "
            "contract are automatically disabled to avoid conflicting changes.",
            "Documents uploaded by vendors now display correctly in the Attached Documents panel "
            "of each policy / security record, with proper file size and download links.",
        ],
    ),
    (
        "Amendments",
        [
            "The Contract Manager's workflow on amendments was reworked. Instead of an "
            "Approve/Reject pair, the manager now sees a clearer Assign Approval button which "
            "opens the approver-selection dialog and routes the amendment through the proper "
            "approval chain. The button is correctly hidden for time-only amendments (which "
            "follow a different path) and for amendments that already have approvers assigned."
        ],
    ),
    (
        "Change Management",
        [
            "The Create Change dialog now offers role-appropriate options:",
            "• Contract Manager can raise Change Directives or Change Orders.",
            "• Vendor / Project Manager can raise Change Requests, Change Orders, and Change "
            "Proposals.",
            "Change Proposals are now correctly scoped to vendors only, and Change Directives are "
            "correctly scoped to managers only — matching standard construction-contract practice.",
        ],
    ),
    (
        "MSA Enhancements",
        [
            "New dialogs for Savings and Holdbacks, with proper visibility for each role.",
            "Invoices gained a new Active status and stricter gating so managers cannot act on "
            "payments that are still awaiting approval.",
            "Server-side filtering by invoice ID for faster searches on large MSAs.",
            "Manager approval and comment flows added to MSA invoices, claims, and changes.",
        ],
    ),
    (
        "Role Coverage",
        [
            "The Project Manager role was introduced where it was missing — most notably across "
            "MSA creation, contract personnel, and the dashboard role-based views. The dashboard "
            "now correctly recognises all current user roles for analytics and navigation."
        ],
    ),
    (
        "Collaboration Tool",
        [
            "Substantial investment in the document collaboration area:",
            "• Offline support — local edits are saved while disconnected and sync when the "
            "connection returns.",
            "• Word-style redline — users can mark inserted or deleted text directly in the "
            "document, with author and timestamp.",
            "• Inline comments — a floating comment button appears beside any text selection "
            "(Google Docs style), letting users pin a discussion to a specific passage.",
            "• @mentions — comments can tag other users attached to the contract; the suggestion "
            "list is restricted to actual contract members.",
            "• Threaded replies — discussions stay organised under the original comment.",
            "• AI Polish (new) — the document scans for all redline edits, sends them to the "
            "configured AI service, and shows professional rephrases the user can approve and "
            "apply with one click, or dismiss.",
        ],
    ),
    (
        "Smaller Polish",
        [
            "Subscription-plan dropdown labels cleaned up so they show only the plan name.",
            "Various small UI fixes: status badges, table layouts, dark-mode contrast, "
            "accessibility labels.",
        ],
    ),
]

for heading, paragraphs in sections:
    h = doc.add_paragraph()
    hr = h.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x2A, 0x44, 0x67)
    for p in paragraphs:
        doc.add_paragraph(p)

# Bottom line
doc.add_paragraph()
bl = doc.add_paragraph()
br = bl.add_run("Bottom line: ")
br.bold = True
bl.add_run(
    "The platform got materially better for three distinct audiences — contract "
    "managers (cleaner controls, fewer chances to act prematurely), vendors "
    "(clearer submission flows, role-appropriate change types), and document "
    "reviewers (real-time collaboration with AI assistance). Multi-currency "
    "support also opens the door for international contracts that the system "
    "previously couldn't represent cleanly."
)

out = r"C:\Users\HomePC\Documents\GitHub\swifter\Product-Update-2026-04-27-to-2026-05-11.docx"
doc.save(out)
print(out)
